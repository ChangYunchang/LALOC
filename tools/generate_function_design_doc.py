from __future__ import annotations

import os
import re
import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DELIVERABLES = ROOT / "deliverables"
ASSETS = DELIVERABLES / "assets"
OUT = DELIVERABLES / "城市低空物流运营中心（LALOC）功能设计说明书V1.0.docx"

IMAGEGEN_DIR = Path.home() / ".codex" / "generated_images" / "019f02ce-35ea-7293-8618-c9351f0e7277"
RAW_HIERARCHY = IMAGEGEN_DIR / "ig_012fde507fccd917016a3e2ae4c85c819aa9b542e3e600a1e3.png"
RAW_ARCH = IMAGEGEN_DIR / "ig_012fde507fccd917016a3e2b1e97bc819a80c09eb4f492fb06.png"

FONT_CN = r"C:\Windows\Fonts\simsun.ttc"
FONT_CN_BOLD = r"C:\Windows\Fonts\simhei.ttf"
FONT_EN = r"C:\Windows\Fonts\times.ttf"


def ensure_dirs() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)


def font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size=size)


def draw_centered(draw: ImageDraw.ImageDraw, box, text: str, fnt, fill=(20, 20, 20), spacing=4, max_chars=12) -> None:
    x1, y1, x2, y2 = box
    lines = []
    for raw in text.split("\n"):
        if len(raw) <= max_chars:
            lines.append(raw)
        else:
            step = max_chars
            lines.extend(raw[i : i + step] for i in range(0, len(raw), step))
    heights = []
    widths = []
    for line in lines:
        b = draw.textbbox((0, 0), line, font=fnt)
        widths.append(b[2] - b[0])
        heights.append(b[3] - b[1])
    total_h = sum(heights) + spacing * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for i, line in enumerate(lines):
        x = x1 + ((x2 - x1) - widths[i]) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += heights[i] + spacing


def draw_box(draw, xy, text, title=False, fill=(248, 248, 248), outline=(40, 40, 40), size=None):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    draw_centered(draw, xy, text, font(FONT_CN_BOLD if title else FONT_CN, size or (34 if title else 30)), max_chars=12)


def make_hierarchy_diagram() -> Path:
    src_copy = ASSETS / "imagegen_hierarchy_background.png"
    if RAW_HIERARCHY.exists():
        shutil.copy2(RAW_HIERARCHY, src_copy)
    img = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(img)

    draw.rectangle((20, 20, 1780, 1180), outline=(70, 70, 70), width=2)
    draw_centered(draw, (530, 45, 1270, 115), "子系统与功能模块层次图", font(FONT_CN_BOLD, 40), max_chars=20)
    draw_box(draw, (675, 155, 1125, 240), "统一低空物流\nGIS 运营平台", True, (235, 235, 235), size=32)

    core_boxes = [
        ((70, 330, 455, 455), "综合态势大屏\n指标｜航线｜回放｜天气"),
        ((515, 330, 900, 455), "智能航路规划\n选点｜巡逻｜应急"),
        ((960, 330, 1345, 455), "安全缓冲区分析\n配置｜仿真｜预警"),
        ((1405, 330, 1730, 455), "安全热力分析\n热力｜热点｜统计"),
    ]
    support_boxes = [
        ((95, 660, 380, 780), "空域管理\n禁飞｜限高｜合规"),
        ((420, 660, 705, 780), "气象保障\n天气｜适飞｜预警"),
        ((745, 660, 1030, 780), "航线与路径服务\n存储｜搜索｜剖面"),
        ((1070, 660, 1355, 780), "安全监管\n冲突｜事件｜台账"),
        ((1395, 660, 1680, 780), "统计与系统支撑\n统计｜数据｜系统"),
    ]

    for xy, text in core_boxes:
        draw.line((900, 240, (xy[0] + xy[2]) // 2, xy[1]), fill=(60, 60, 60), width=2)
        draw_box(draw, xy, text, False)
    draw.line((900, 455, 900, 600), fill=(60, 60, 60), width=2)
    draw_centered(draw, (650, 565, 1150, 635), "共享数据、接口服务与业务规则支撑", font(FONT_CN_BOLD, 32), max_chars=16)
    for xy, text in support_boxes:
        draw.line((900, 635, (xy[0] + xy[2]) // 2, xy[1]), fill=(90, 90, 90), width=2)
        draw_box(draw, xy, text, False, (252, 252, 252), (80, 80, 80))

    notes = [
        "四个核心子系统对应当前前端 GIS 业务入口，支撑功能通过后端 API、数据库和基础数据资源提供服务。",
        "图中功能模块按职责边界组织：展示、规划、仿真分析、风险统计与支撑治理分离。"
    ]
    y = 930
    for n in notes:
        draw.text((95, y), n, font=font(FONT_CN, 28), fill=(35, 35, 35))
        y += 48

    out = ASSETS / "laloc_subsystem_hierarchy.png"
    img.save(out, quality=95)
    return out


def make_arch_diagram() -> Path:
    src_copy = ASSETS / "imagegen_architecture_background.png"
    if RAW_ARCH.exists():
        shutil.copy2(RAW_ARCH, src_copy)
    img = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((20, 20, 1780, 1180), outline=(70, 70, 70), width=2)
    draw_centered(draw, (500, 35, 1300, 100), "系统总体架构图", font(FONT_CN_BOLD, 42))

    layers = [
        ("用户与访问层", "低空管理部门｜物流运营人员｜系统管理员"),
        ("前端表现层", "页面框架｜组件库｜状态管理｜路由｜统计图表"),
        ("GIS 地图层", "统一地图接口｜高德二维地图｜三维实景地球｜图层图例"),
        ("业务服务层", "接口服务｜空域｜天气｜路径｜航线｜安全｜统计"),
        ("算法与规则层", "路径搜索｜禁飞缓冲｜建筑避让｜安全距离｜热力采样"),
        ("数据存储层", "空间数据库｜航线｜空域｜建筑｜任务｜事件"),
        ("外部数据层", "地图服务｜天气服务｜建筑白模｜空间文件｜建筑缓存"),
    ]
    top = 130
    for idx, (name, desc) in enumerate(layers):
        y1 = top + idx * 135
        y2 = y1 + 92
        draw.rounded_rectangle((95, y1, 360, y2), radius=10, fill=(232, 232, 232), outline=(45, 45, 45), width=2)
        draw_centered(draw, (95, y1, 360, y2), name, font(FONT_CN_BOLD, 30), max_chars=7)
        draw.rounded_rectangle((400, y1, 1645, y2), radius=10, fill=(250, 250, 250), outline=(70, 70, 70), width=2)
        draw_centered(draw, (400, y1, 1645, y2), desc, font(FONT_CN, 30), max_chars=32)
        if idx < len(layers) - 1:
            draw.line((1015, y2, 1015, y2 + 36), fill=(55, 55, 55), width=3)
            draw.polygon([(1015, y2 + 46), (1005, y2 + 30), (1025, y2 + 30)], fill=(55, 55, 55))

    draw.text((106, 1115), "说明：前端可通过 Vite Mock API 独立调试核心 GIS 功能；完整业务能力由后端服务与 PostGIS 数据库支撑。", font=font(FONT_CN, 26), fill=(35, 35, 35))
    out = ASSETS / "laloc_system_architecture.png"
    img.save(out, quality=95)
    return out


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="999999", size="6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, align=None, first_line=True):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if first_line and style is None and text:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    for r in p.runs:
        set_run_font(r, 14 if level == 1 else 12 if level == 2 else 11, True, "000000")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "F2F2F2")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(p.text) < 12 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.2
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    set_run_font(r, font_size, row == table.rows[0])
            if widths:
                cell.width = Cm(widths[i])
    doc.add_paragraph()
    return table


def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.line_spacing = 1.5

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("城市低空物流运营中心（LALOC）功能设计说明书")
    set_run_font(run, 9)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("第 ")
    set_run_font(r, 9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    rr = OxmlElement("w:r")
    tt = OxmlElement("w:t")
    tt.text = "1"
    rr.append(tt)
    fld.append(rr)
    footer._p.append(fld)
    r2 = footer.add_run(" 页")
    set_run_font(r2, 9)
    return doc


core_modules = [
    {
        "code": "LALOC-GIS-01",
        "name": "综合态势大屏",
        "route": "/dashboard",
        "summary": "综合态势大屏汇聚低空物流航线、无人机、空域限制、天气和运行指标，通过二维和三维地图形成统一态势视图，是管理人员进入系统后的首要监控入口。",
        "sub": [
            ("核心指标卡片", "展示禁飞区、限高区、航线数量、运行无人机等核心指标。", "页面加载或数据刷新时触发。", "调用空域统计、航线列表和天气接口，聚合后以卡片形式展示。", "指标卡采用紧凑布局，支持快速识别系统运行规模和限制区域数量。", "当接口失败时保留占位值并提示数据加载异常。"),
            ("航线列表与筛选", "展示模拟航线、接口航线和本地保存航线，支持用户按名称、距离、时间和状态筛选。", "用户在航线面板输入筛选条件或点击列表项。", "系统合并样例航线与接口返回数据，维护统一航线对象结构。", "列表展示航线名称、距离、预计时间和状态，选中项与地图联动。", "无匹配航线时显示空状态，不影响地图基础图层。"),
            ("航线高亮", "突出当前关注航线，弱化其他航线，辅助用户定位任务走廊。", "用户点击航线列表项或地图航线对象。", "调用 MapContainer 的高亮接口，将选中航线渲染为强调线型。", "选中航线使用深色实线，其余航线降低透明度。", "未获取地图实例时延迟执行高亮请求。"),
            ("无人机动画", "模拟无人机沿航线循环飞行，表现飞行阶段和高度变化。", "页面进入、切换航线或点击播放控制时触发。", "根据航线 path 和高度剖面计算当前位置、朝向和阶段。", "2D 模式显示动态图标，3D 模式显示贴地高度转换后的空间位置。", "路径点不足时不启动动画并给出提示。"),
            ("时间轴回放", "按时间进度复现无人机历史位置，用于过程检查和事件复盘。", "用户切换回放模式、拖动时间轴或调整倍速。", "系统按照 progress 值插值航线坐标，更新无人机位置和状态。", "时间轴显示当前进度、播放状态和 0.5 到 5 倍速控制。", "未选择航线时禁用回放控件。"),
            ("3D 视角跟踪", "在三维场景中自动跟随目标无人机，提高空间态势感知能力。", "用户在 3D 回放模式下启用跟随。", "Cesium 相机根据无人机位置、高度和航向动态设置视角。", "镜头保持适度俯仰角，保证航线和周边建筑同时可见。", "地形采样超时时回退到椭球体高度。"),
            ("天气面板", "展示实时温度、湿度、风向风力和飞行适宜性。", "页面加载、定时刷新或用户手动刷新。", "调用 weather/live 与 weather/flyable 接口，并根据阈值生成适飞结论。", "面板以文本和状态标识展示天气条件，不采用彩色复杂图表。", "天气接口不可用时使用回退数据并标记更新时间。"),
            ("图层图例", "说明禁飞区、限高区、航线阶段等地图符号含义。", "用户展开图例或切换图层显示。", "系统读取当前地图图层状态，控制图例项显示。", "图例与地图状态保持一致，避免用户误读空间对象。", "图层未加载时显示加载中状态。"),
        ],
    },
    {
        "code": "LALOC-GIS-02",
        "name": "智能航路规划",
        "route": "/path-planning、/emergency-routing",
        "summary": "智能航路规划子系统面向配送、巡逻和应急场景，综合禁飞区、限高区、建筑、天气和无人机参数生成安全、可执行的航路。",
        "sub": [
            ("地图点击选点", "允许用户在 2D 或 3D 地图上直接指定起点、终点和途经点。", "用户点击地图或切换选点模式。", "MapContainer 将点击坐标转换为经纬度并写入规划表单。", "地图上显示不同形态的起点、终点和途经点标记。", "点位落入禁飞区时拒绝写入并提示原因。"),
            ("坐标输入", "支持用户手动录入 lng,lat 坐标，适配精确规划需求。", "用户在输入框提交坐标文本。", "系统解析经纬度格式，校验数值范围和点位合法性。", "合法坐标同步到地图标记，非法坐标显示输入错误。", "空值、格式错误或超出广州范围时不触发规划。"),
            ("途经点管理", "支持多个途经点按顺序参与分段规划。", "用户添加、删除或调整途经点顺序。", "系统维护 waypoint 数组，并按起点到途经点再到终点拆分规划段。", "途经点列表与地图标记双向联动。", "重复点或距离过近时提示用户调整。"),
            ("禁飞区选点拦截", "避免用户将起降点或途经点设置在禁飞区域内。", "点位选择、坐标输入或规划前校验时触发。", "调用区域判断逻辑或本地 GeoJSON 空间关系判断。", "页面弹窗说明被拦截点位和相关限制区域。", "禁飞区数据未加载时要求先加载空域图层。"),
            ("规划参数配置", "配置无人机速度、巡航高度、安全余量和避障开关。", "用户调整滑条、输入框或复选开关。", "系统将参数组装为 pathfinding 请求或 Mock 规划参数。", "界面显示参数即时值和建议范围。", "超出上下限时自动修正或阻止提交。"),
            ("Theta* 路径搜索", "在栅格空间中搜索任意角度可行路径。", "用户点击开始规划后触发。", "系统将禁飞区缓冲、建筑障碍和限高约束转换为搜索成本。", "规划过程在 3D 模式可显示计算进度。", "无可行路径时返回失败原因并保留用户输入。"),
            ("路径平滑与重采样", "消除栅格路径锯齿，使航线适合无人机连续飞行。", "路径搜索成功后自动执行。", "依次执行视线串拉、Chaikin 角切割和弧长均匀重采样。", "地图上展示平滑曲线而非原始栅格折线。", "平滑后若穿越硬约束，回退到安全折线段。"),
            ("高度剖面生成", "根据起降、巡航、限高和建筑避让需求生成高度序列。", "路径几何确定后触发。", "系统计算每个路径点的 AGL 高度，必要时进行地形采样转换。", "二维以阶段颜色表达高度变化，三维以空间曲线展示。", "地形或建筑数据缺失时使用保守默认高度。"),
            ("阶段着色", "将爬升、巡航、下降、限高绕行和建筑避让分段展示。", "高度剖面生成后自动触发。", "系统为路径点写入阶段字段并映射到渲染样式。", "图例同步解释各阶段含义。", "阶段计算不完整时以普通航线样式展示。"),
            ("航线保存", "将规划结果命名后保存到态势大屏航线列表。", "用户点击保存航线并填写名称。", "系统保存路径点、距离、时间、高度剖面和阶段信息。", "保存成功后可跨页面在大屏中加载。", "名称为空或路径结果失效时阻止保存。"),
            ("沿线飞行巡逻", "支持用户手绘折线并生成沿线巡逻航线。", "用户进入巡逻模式并绘制至少两个点。", "系统按绘制线生成航路点并应用离地高度和速度参数。", "地图实时显示已绘制点数和生成结果。", "绘制点不足时不允许开始规划。"),
            ("空域巡回巡逻", "支持在多边形区域内生成边界巡逻或犁地式覆盖路线。", "用户绘制区域并选择巡逻模式。", "系统按边界采样或 Boustrophedon 扫描生成巡逻路线。", "页面显示路线点数、条带间距和巡逻方向。", "区域自交或面积过小时提示重新绘制。"),
            ("应急告警选择", "按低电量、设备故障和通信中断筛选应急对象。", "用户进入应急页面或选择告警原因。", "系统过滤飞行中无人机列表并标记低电量对象。", "低电量对象在卡片上突出显示。", "无匹配无人机时显示空状态。"),
            ("安全迫降点推荐", "为异常无人机搜索最近可用充电站或维修站。", "用户点击无人机卡片触发告警。", "系统按距离、可用性和续航半径计算推荐排序。", "地图绘制续航圈和候选安全点。", "电量不足到达任何点时提示无法保障迫降。"),
            ("应急航路导出", "输出应急路径 JSON，便于后续处置或记录。", "用户生成应急航路后点击导出。", "系统整理无人机、起点、安全点、路径点和评估结果。", "浏览器下载结构化 JSON 文件。", "无有效应急路径时导出按钮不可用。"),
        ],
    },
    {
        "code": "LALOC-GIS-03",
        "name": "安全缓冲区分析",
        "route": "/safety-buffer/analysis",
        "summary": "安全缓冲区分析子系统通过多航线、多无人机仿真，实时判断水平和垂直安全间隔，识别潜在冲突并给出处置建议。",
        "sub": [
            ("安全范围配置", "配置水平缓冲区、警戒距离和垂直缓冲区。", "用户拖动滑条、输入数值或恢复默认。", "系统校验参数范围并写入仿真判定配置。", "配置面板同步显示当前阈值。", "警戒距离小于缓冲区时提示参数不合理。"),
            ("多航线同时仿真", "模拟多条交叉航线在不同时段的飞行活动。", "用户选择时段并启动播放。", "系统根据时段生成无人机数量和初始位置。", "地图同时显示多条航线和多个无人机对象。", "仿真对象过多时限制最大数量以保持性能。"),
            ("实时碰撞检测", "按时间步计算无人机间三维距离。", "仿真播放、拖动进度或参数变化时触发。", "系统比较水平距离、垂直距离和安全阈值。", "发生风险时高亮相关无人机和缓冲区。", "缺少高度数据时使用水平距离进行降级判断。"),
            ("三级预警判定", "将运行状态划分为正常、警告和冲突。", "每次碰撞检测后自动判定。", "系统根据警戒距离和缓冲区阈值生成等级。", "状态面板显示正常、减速或停止标签。", "等级切换时避免重复刷屏记录。"),
            ("让行策略", "为发生冲突的无人机生成停止或减速建议。", "检测到警告或冲突时触发。", "同航线跟随者优先让行，跨航线按索引或优先级处理。", "事件日志记录建议动作。", "策略仅作为模拟建议，不直接控制真实无人机。"),
            ("无人机状态面板", "展示每架无人机的名称、类型、进度、高度和速度。", "仿真初始化和每个时间步刷新。", "系统从仿真状态对象提取展示字段。", "面板按状态排序，便于优先查看异常对象。", "无人机对象消失时自动移出列表。"),
            ("碰撞事件日志", "记录缓冲区重叠、警告和冲突事件。", "风险等级变化时写入。", "系统生成时间戳、涉及对象、类型和建议。", "日志列表支持滚动查看。", "重复事件合并，避免同一时刻产生大量冗余记录。"),
            ("时间轴与速度控制", "控制仿真进度、播放状态和倍速。", "用户点击播放、暂停、重置或调整倍速。", "系统按倍速推进仿真时间并更新地图对象。", "滑条显示当前时间片进度。", "进度到达终点后自动停止或循环。"),
            ("2D/3D 缓冲区渲染", "在二维和三维地图中展示安全缓冲范围。", "地图模式切换或无人机位置更新时触发。", "2D 绘制圆形缓冲区，3D 绘制圆柱体空间范围。", "半透明灰阶样式避免遮挡底图。", "3D 资源未加载时回退到 2D 表达。"),
        ],
    },
    {
        "code": "LALOC-GIS-04",
        "name": "安全热力分析",
        "route": "/density/contour、/density/hotspot、/density/stats",
        "summary": "安全热力分析子系统聚焦低空密度和拥堵风险，通过热力图、热点识别和统计图表辅助管理人员发现重点空域。",
        "sub": [
            ("低空密度采样", "沿核心航线走廊生成密度采样点。", "页面加载或时段切换时触发。", "系统结合 9 条模拟航线、随机抖动和扩散点生成样本。", "采样点用于热力图和统计分析，不直接展示全部点位。", "采样为空时显示无数据提示。"),
            ("安全风险热力图", "以热力图展示低空运行密度和风险空间分布。", "用户进入热力分析页面或调整参数。", "系统将采样点转换为高德 HeatMap 数据结构。", "图层以灰阶深浅表达密度强弱，符合文档黑白规范。", "地图插件不可用时展示降级提示。"),
            ("时间段筛选", "支持早高峰、午间、下午和晚高峰分析。", "用户点击时段选项。", "系统按时段因子重新计算密度权重。", "当前时段信息在面板中显示。", "未选择时段时默认使用全时段。"),
            ("时间轴动画", "按 5 分钟时间片动态展示密度变化。", "用户点击播放、暂停或重置。", "系统随时间片更新热力数据和信息牌。", "页面显示当前时刻和采样点数量。", "播放期间切换页面时停止计时器。"),
            ("透明度控制", "调整热力图覆盖层可见程度。", "用户拖动透明度滑条。", "系统更新热力图图层 opacity。", "底图、航线和热力层保持可辨识。", "透明度过低时提示可能影响观察。"),
            ("拥堵热点识别", "根据密度阈值和半径识别低空拥堵区域。", "用户进入热点页面或修改阈值。", "系统遍历预置热点和密度统计结果，生成高、中、低等级。", "地图显示圆形热点标记，列表显示坐标和架次。", "阈值设置过高时可能无热点，页面给出说明。"),
            ("热点列表联动", "点击列表项后地图聚焦到对应热点。", "用户点击热点列表记录。", "系统调用地图定位并打开信息窗。", "信息窗展示密度值、峰值时段和关联航线。", "热点坐标异常时不执行地图跳转。"),
            ("区域密度统计卡片", "展示总飞行架次、高密度空域数和预警空域数。", "统计页面加载或筛选条件变化时触发。", "系统汇总密度数据并计算关键指标。", "卡片布局便于管理人员快速读数。", "统计数据缺失时显示 0 和无数据说明。"),
            ("高密度空域排名", "以柱状图展示重点区域密度排序。", "用户选择 Top 数量或时段。", "系统按密度值排序并截取前 N 个区域。", "图表显示区域名称和密度等级。", "区域名称过长时自动换行或缩略。"),
            ("高频航线统计", "统计核心航线的使用频次和排名。", "用户切换统计维度。", "系统聚合航线架次，生成水平柱状图数据。", "图表适合比较航线负荷。", "航线数据为空时提示先加载航线。"),
            ("24 小时趋势分析", "展示全天密度变化趋势、峰值和均值。", "页面加载或筛选变化时触发。", "系统按小时聚合密度值并计算均值线。", "折线图标注峰值时段。", "缺少连续时间数据时使用模拟序列补齐。"),
        ],
    },
]


support_modules = [
    ("SUP-01", "空域管理支撑", "维护禁飞区、限高区和空间查询能力，为规划、态势展示和合规审查提供统一规则。", ["禁飞区查询", "限高区查询", "点位约束检查", "区域统计", "航线合规审查"]),
    ("SUP-02", "气象保障支撑", "接入实时天气和飞行适宜性判断，支持大屏展示、规划约束和应急分析。", ["实时天气获取", "天气缓存更新", "适飞判断", "气象异常提示", "外部服务回退"]),
    ("SUP-03", "航线与路径服务支撑", "统一管理航线存储、路径规划结果、高度剖面和跨页面共享数据。", ["航线列表", "航线详情", "航线创建", "路径请求校验", "高度剖面管理"]),
    ("SUP-04", "物流运营支撑", "面向企业、订单、任务、站点和无人机资源管理，为真实运营闭环预留业务能力。", ["企业档案", "订单管理", "任务管理", "站点管理", "无人机资源"]),
    ("SUP-05", "安全监管支撑", "提供冲突检测、拥堵识别、风险热力、异常事件和监管台账能力。", ["冲突检查", "拥堵分析", "风险热力", "事件处置", "安全记录归档"]),
    ("SUP-06", "统计决策支撑", "对城市运行、企业效率、服务质量、成本和站点布局进行统计分析。", ["城市概览", "任务趋势", "航线利用", "企业效率", "站点布局"]),
    ("SUP-07", "数据资源支撑", "管理 Shapefile、GeoJSON、OSM 建筑、基础地理数据和外部数据服务。", ["空间数据导入", "三维建筑缓存", "GIS 图层", "数据质检", "外部数据源"]),
    ("SUP-08", "系统管理支撑", "维护用户、参数、日志、GIS 图层和服务状态，保障平台安全稳定运行。", ["用户管理", "参数配置", "操作日志", "服务状态", "图层管理"]),
]


def add_module_design(doc, module, idx):
    add_heading(doc, f"3.{idx} {module['name']}功能模块设计（{module['code']}）", 2)
    add_heading(doc, "概述", 3)
    add_para(doc, module["summary"])
    add_para(doc, f"该模块对应系统入口 {module['route']}，主要服务于低空管理部门对城市低空物流运行状态的持续观察、分析和处置。模块设计强调地图联动、数据一致、操作反馈明确和异常状态可解释。")
    add_heading(doc, "功能目标", 3)
    for goal in [
        "形成清晰的业务入口和操作路径，使用户能够在地图、面板和图表之间快速理解业务状态。",
        "将空间数据、业务数据和分析结果组织为统一对象，便于跨页面共享和后续扩展。",
        "在关键操作处提供校验、提示和异常回退，避免错误数据继续进入业务流程。",
    ]:
        add_bullet(doc, goal)
    add_heading(doc, "子功能清单", 3)
    add_table(doc, ["序号", "子功能", "用途"], [[f"{module['code']}-{i:02d}", s[0], s[1]] for i, s in enumerate(module["sub"], 1)], [2.4, 4.0, 9.0], 9)
    add_heading(doc, "子功能详细设计", 3)
    for i, (name, purpose, trigger, logic, display, exception) in enumerate(module["sub"], 1):
        add_heading(doc, f"{module['code']}-{i:02d} {name}", 3)
        add_para(doc, f"子功能用途：{purpose}该子功能是“{module['name']}”模块中的独立功能点，应在界面、数据处理和异常提示上保持独立边界，避免与其他子功能混合实现。")
        add_para(doc, f"用户操作或触发方式：{trigger}触发后系统应立即给出可见反馈，确保用户能够判断当前操作是否已被接收。")
        add_para(doc, f"系统处理逻辑：{logic}处理过程中应复用统一的接口封装、地图协调器或业务状态对象，保证 2D/3D 展示和后端数据含义一致。")
        add_para(doc, f"页面、地图或图表表现：{display}展示时应遵循浅色、克制、专业的界面风格，避免使用与业务含义无关的装饰性视觉元素。")
        add_para(doc, f"使用的数据或接口：该子功能主要使用 README 中描述的航线、空域、天气、路径规划、热力采样、仿真状态或统计数据。接口层应优先通过 `frontend/src/api` 中的封装调用，地图展示通过 `MapContainer` 或对应地图视图完成。")
        add_para(doc, f"异常提示或边界限制：{exception}异常发生后应保留用户已输入信息，不应清空有效上下文；对于无法继续执行的操作，应明确说明原因和下一步处理方式。")
    add_heading(doc, "输入数据", 3)
    add_para(doc, "输入数据包括用户交互参数、地图点位、航线对象、空域 GeoJSON、天气数据、无人机状态、路径规划请求和统计筛选条件等。所有输入在进入核心计算前都应进行格式、范围和业务合法性校验。")
    add_heading(doc, "输出结果", 3)
    add_para(doc, "输出结果包括地图图层、航线几何、无人机状态、指标卡片、分析图表、事件日志、导出文件和用户提示信息。输出内容应可被用户理解，并能被其他页面或支撑服务复用。")
    add_heading(doc, "处理流程", 3)
    for step in ["用户进入模块并加载基础图层和业务数据。", "用户通过地图、表单或控制组件提交操作。", "系统执行参数校验、接口调用、空间计算或仿真分析。", "系统刷新地图对象、面板状态和结果说明。", "发生异常时保留上下文并给出可操作提示。"]:
        add_numbered(doc, step)
    add_heading(doc, "异常情况", 3)
    for err in ["地图服务或 3D 资源加载失败时，应提供 2D 或静态数据回退。", "接口超时或返回异常时，应提示用户并避免使用半成品结果。", "用户输入不完整、不合法或与空域约束冲突时，应阻止后续计算。"]:
        add_bullet(doc, err)
    add_heading(doc, "与其他模块关系", 3)
    add_para(doc, f"{module['name']}与空域管理、气象保障、航线与路径服务、安全监管和统计决策支撑能力存在数据交互关系。模块自身负责业务展示和交互闭环，不重复实现支撑模块的数据维护职责。")


def build_document():
    ensure_dirs()
    hierarchy = make_hierarchy_diagram()
    arch = make_arch_diagram()

    doc = setup_doc()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("城市低空物流运营中心（LALOC）\n功能设计说明书")
    set_run_font(r, 22, True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("版本：V1.0\n发布日期：" + date.today().isoformat())
    set_run_font(r2, 12)
    doc.add_page_break()

    add_table(doc, ["版本", "版本说明", "修改人", "审核人", "批准人", "发布日期"], [["V1.0", "形成城市低空物流运营中心功能设计说明书初稿", "", "", "", date.today().isoformat()]], [2.0, 6.2, 2.0, 2.0, 2.0, 3.0], 9)
    doc.add_page_break()

    add_heading(doc, "1 引言", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本文档用于说明城市低空物流运营中心（LALOC）的功能设计方案，明确系统建设目标、子系统划分、功能模块构成、子功能设计、输入输出、处理流程和非功能需求。文档既面向当前已实现的 GIS 核心功能，也为后续支撑能力扩展提供结构化依据。")
    add_para(doc, "通过本说明书，项目开发、测试、验收和维护人员能够理解系统各部分的功能边界，避免将展示、规划、仿真、分析和支撑管理职责混淆。")
    add_heading(doc, "1.2 读者对象", 2)
    for reader in ["项目指导教师、评审人员和验收人员。", "前端、后端、GIS、数据库和测试开发人员。", "低空管理部门业务人员、物流运营企业代表和系统管理员。"]:
        add_bullet(doc, reader)
    add_heading(doc, "1.3 参考资料", 2)
    for ref in ["项目 README.md。", "项目 Structure.md。", "功能设计说明书 V1.0 模板。", "前端 Vue 3、Vite、Element Plus、Pinia、ECharts 相关实现。", "后端 FastAPI、PostgreSQL、PostGIS、SQLAlchemy 和 GeoAlchemy2 相关实现。"]:
        add_bullet(doc, ref)

    add_heading(doc, "2 项目概述", 1)
    add_heading(doc, "2.1 项目背景与建设目标", 2)
    add_para(doc, "随着无人机物流行业快速发展，城市低空空域管理需要同时面对飞行活动增长、空域限制复杂、运行安全要求高和监管可视化不足等问题。LALOC 面向低空管理部门，以 2D/3D GIS 地图为核心载体，提供低空物流运行态势监控、智能航路规划、安全缓冲区分析和低空密度热力分析能力。")
    add_para(doc, "系统建设目标是形成一套可展示、可规划、可仿真、可分析、可支撑扩展的低空物流运营管理平台。当前项目以四个 GIS 核心子系统作为主要业务入口，同时通过后端接口、PostGIS 数据库、空域数据、气象服务和系统管理能力支撑完整业务闭环。")
    add_heading(doc, "2.2 用户角色与使用场景", 2)
    add_table(doc, ["角色", "主要关注点", "典型场景"], [
        ["低空管理人员", "城市运行态势、安全风险、拥堵热点和异常处置", "查看大屏、分析热力、评估冲突、复盘航线"],
        ["物流运营人员", "航路生成、任务可执行性、无人机状态和应急处置", "选点规划、巡逻路线生成、应急迫降分析"],
        ["系统管理员", "数据、参数、服务状态和用户权限", "维护系统参数、检查服务、管理图层和日志"],
    ], [3.0, 6.0, 6.0])
    add_heading(doc, "2.3 总体功能结构", 2)
    add_para(doc, "系统采用“4 个 GIS 核心子系统 + 支撑功能模块”的结构。四个核心子系统直接面向用户操作，支撑功能模块提供数据、接口、规则和管理能力。该划分既贴合当前 README 描述的功能入口，又为后端已有的空域、物流、安全、统计和系统管理接口留下清晰扩展边界。")
    add_table(doc, ["编号", "子系统或模块", "职责说明"], [[m["code"], m["name"], m["summary"]] for m in core_modules] + [[c, n, d] for c, n, d, _ in support_modules], [3.0, 4.0, 8.0])
    add_heading(doc, "2.4 四个 GIS 核心子系统划分", 2)
    for m in core_modules:
        add_para(doc, f"{m['name']}：{m['summary']}入口路由为 {m['route']}。")
    add_heading(doc, "2.5 支撑功能模块划分", 2)
    for code, name, desc, subs in support_modules:
        add_para(doc, f"{name}（{code}）：{desc}主要包含{ '、'.join(subs) }等支撑能力。")
    add_heading(doc, "2.6 子系统与功能模块层次图", 2)
    doc.add_picture(str(hierarchy), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "图 2-1 子系统与功能模块层次图。该图由 imagegen 生成专业灰阶底稿，并通过本地确定性排版叠加中文标签。", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_heading(doc, "2.7 系统总体架构图", 2)
    doc.add_picture(str(arch), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "图 2-2 系统总体架构图。系统由用户访问层、前端表现层、GIS 地图层、业务服务层、算法规则层、数据存储层和外部数据层构成。", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_heading(doc, "2.8 前后端、GIS、数据、接口协同关系", 2)
    add_para(doc, "前端采用 Vue 3 + Vite 构建，利用 MapContainer 协调高德 2D 地图和 Cesium 3D 地图，屏蔽不同地图引擎差异。后端采用 FastAPI 提供 RESTful API，围绕 zones、weather、pathfinding、routes、airspace、logistics、safety、statistics 和 system 等模块组织服务。数据库采用 PostgreSQL + PostGIS 保存空域、建筑、航线、物流、事件和系统数据。")
    add_para(doc, "在前端独立开发模式下，Vite Mock API 能够支撑禁飞区、限高区、天气、航线和路径规划等核心 GIS 功能调试；在完整后端模式下，系统通过 API 与数据库提供更完整的数据管理、监管统计和系统管理能力。")

    add_heading(doc, "3 功能设计说明", 1)
    add_heading(doc, "3.1 总体功能分析", 2)
    add_para(doc, "功能设计按照子系统、功能模块、子功能三个层次展开。每个子功能单独说明用途、触发方式、处理逻辑、表现形式、使用数据和异常边界，保证开发人员能够按独立功能点实现和测试。")
    add_table(doc, ["模块编号", "功能模块", "说明"], [[m["code"], m["name"], m["summary"]] for m in core_modules] + [[c, n, d] for c, n, d, _ in support_modules], [3.0, 4.0, 8.0])
    for i, module in enumerate(core_modules, 2):
        add_module_design(doc, module, i)

    add_heading(doc, "3.6 支撑功能模块设计（LALOC-SUP）", 2)
    add_heading(doc, "概述", 3)
    add_para(doc, "支撑功能模块不作为当前前端四个 GIS 入口中的独立页面展开，但它们为核心子系统提供数据、规则、接口和管理能力。支撑功能设计重点描述能力边界、接口协同和对子系统的服务关系。")
    add_table(doc, ["支撑编号", "支撑模块", "子功能"], [[c, n, "、".join(subs)] for c, n, _, subs in support_modules], [3.0, 4.0, 8.0])
    for code, name, desc, subs in support_modules:
        add_heading(doc, f"{code} {name}", 3)
        add_para(doc, f"模块概述：{desc}该支撑能力通过后端接口、数据库表、GIS 图层或配置参数为四个核心 GIS 子系统提供基础服务。")
        for j, sub in enumerate(subs, 1):
            add_para(doc, f"{code}-{j:02d} {sub}：该子功能负责{name}中的“{sub}”能力。系统应明确输入数据来源、输出结果形式和调用边界；当核心子系统请求该能力时，支撑模块返回标准化数据，不直接承担前端页面展示职责。")
            add_para(doc, f"处理要求：{sub}需要进行基础参数校验、数据存在性检查和异常返回封装。接口不可用、数据为空或业务规则不满足时，应返回清晰错误信息，便于前端模块展示用户可理解的提示。")
        add_para(doc, f"与核心模块关系：{name}主要服务于综合态势大屏、智能航路规划、安全缓冲区分析和安全热力分析。核心模块使用其结果，不重复维护支撑数据本身。")

    add_heading(doc, "4 系统非功能设计", 1)
    nonfunc = [
        ("4.1 数据需求", "系统数据包括禁飞区、限高区、航线、建筑、天气、无人机、任务、站点、事件、统计结果、用户、参数、日志和 GIS 图层。空间数据应支持 GeoJSON、Shapefile 和 PostGIS 几何对象，业务数据应满足后端 Pydantic 模型校验要求。"),
        ("4.2 环境需求", "前端运行环境为 Node.js 20+、Vue 3、Vite 5；后端运行环境为 Python 3.12+、FastAPI、Uvicorn；数据库运行环境为 PostgreSQL 16 + PostGIS 3.4，并通过 Docker Compose 提供本地部署能力。"),
        ("4.3 性能需求", "核心 GIS 页面应在常规浏览器中保持流畅交互。二维地图、热力图和图表应避免一次性渲染过多对象；三维场景应采用懒加载策略，Cesium 资源仅在用户切换 3D 模式时加载。"),
        ("4.4 安全需求", "系统应区分管理部门、企业运营人员和系统管理员的访问权限。后端接口应进行输入校验，系统参数、用户、日志和企业接入信息应由系统管理能力统一维护。"),
        ("4.5 可用性与兼容性需求", "系统应支持前端 Mock 模式和完整后端模式。外部地图、气象或三维资源不可用时，应提供降级提示或回退数据，避免核心页面完全不可用。"),
        ("4.6 可维护性需求", "前端新增页面应优先通过 MapContainer 使用地图能力，避免重复创建地图实例。后端新增业务模型应在 models、database 和 alembic 环境中统一注册，保证迁移和 API 扩展一致。"),
        ("4.7 其它", "文档、接口、数据和图层命名应保持一致。所有功能模块应保留异常提示和日志记录能力，便于后续测试、演示和问题定位。"),
    ]
    for title, body in nonfunc:
        add_heading(doc, title, 2)
        add_para(doc, body)
        for extra in ["应明确数据来源、更新周期和异常回退方式。", "应保证用户操作有反馈，避免长时间无响应。", "应保持模块职责清晰，支撑功能与展示功能不互相侵入。"]:
            add_bullet(doc, extra)

    add_heading(doc, "5 附件", 1)
    add_heading(doc, "5.1 功能模块编号表", 2)
    rows = []
    for m in core_modules:
        for i, s in enumerate(m["sub"], 1):
            rows.append([f"{m['code']}-{i:02d}", m["name"], s[0]])
    add_table(doc, ["内容编号", "所属模块", "建设内容"], rows, [4.0, 4.0, 7.0], 8.5)
    add_heading(doc, "5.2 术语说明", 2)
    add_table(doc, ["术语", "说明"], [
        ["LALOC", "Low-Altitude Logistics Operations Center，城市低空物流运营中心。"],
        ["GIS", "Geographic Information System，地理信息系统。"],
        ["AGL", "Above Ground Level，地面以上高度。"],
        ["PostGIS", "PostgreSQL 的空间数据库扩展，用于空间查询和几何对象管理。"],
        ["Theta*", "基于 A* 的任意角度路径搜索算法，用于生成更平滑的可行航路。"],
    ], [4.0, 11.0])
    add_heading(doc, "5.3 参考接口与路由清单", 2)
    add_table(doc, ["类别", "接口或路由", "说明"], [
        ["前端路由", "/dashboard", "综合态势大屏。"],
        ["前端路由", "/path-planning", "智能路径规划与巡逻路线生成。"],
        ["前端路由", "/emergency-routing", "应急航路规划。"],
        ["前端路由", "/safety-buffer/analysis", "安全缓冲区分析。"],
        ["前端路由", "/density/contour", "安全风险热力分析。"],
        ["前端路由", "/density/hotspot", "低空拥堵识别。"],
        ["前端路由", "/density/stats", "区域密度统计。"],
        ["后端接口", "/api/zones", "禁飞区、限高区、统计和点位检查。"],
        ["后端接口", "/api/weather", "实时天气、预报和适飞判断。"],
        ["后端接口", "/api/pathfinding/plan", "路径规划请求。"],
        ["后端接口", "/api/routes", "航线管理。"],
        ["后端接口", "/api/safety", "冲突、拥堵、风险热力和事件台账。"],
        ["后端接口", "/api/statistics", "统计决策分析。"],
        ["后端接口", "/api/system", "用户、参数、日志、服务状态和 GIS 图层。"],
    ], [3.2, 5.2, 6.6], 8.8)

    doc.core_properties.title = "城市低空物流运营中心（LALOC）功能设计说明书"
    doc.core_properties.subject = "功能设计说明书"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_document()
    print(path)
